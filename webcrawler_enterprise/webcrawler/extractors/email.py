"""Email address extraction from text and documents."""

from __future__ import annotations

import html
import re
from pathlib import Path

EMAIL_RE = re.compile(
    r"(?<![A-Za-z0-9._%+-])"
    r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
    r"(?![A-Za-z0-9._%+-])",
    re.IGNORECASE,
)

# info [at] qau [dot] edu [dot] pk  /  info(at)qau(dot)edu(dot)pk
# Intentionally excludes plain "@" so normal emails are handled by EMAIL_RE.
OBFUSCATED_EMAIL_RE = re.compile(
    r"(?P<local>[A-Za-z0-9._%+-]+)\s*(?:\[at\]|\(at\)|\bat\b|&#64;|&#x40;)\s*"
    r"(?P<domain>[A-Za-z0-9-]+(?:\s*(?:\[dot\]|\(dot\)|\bdot\b|\.)\s*[A-Za-z0-9-]+)+)",
    re.IGNORECASE,
)

CFEMAIL_RE = re.compile(
    r"data-cfemail=[\"']([0-9a-fA-F]+)[\"']|/cdn-cgi/l/email-protection#([0-9a-fA-F]+)",
    re.IGNORECASE,
)

INVALID_SUFFIXES = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".svg",
    ".webp",
    ".css",
    ".js",
    ".woff",
    ".woff2",
}
INVALID_LOCAL_PARTS = {"email", "name", "user", "username", "example", "domain", "your"}


# Common false positives pulled from PDF hyphenation / layout noise.
INVALID_TLDS = {
    "https",
    "http",
    "also",
    "moreover",
    "modern",
    "ten",
    "pdf",
    "docx",
    "html",
    "com0",
    "edu0",
}


def _clean_email(email: str) -> str | None:
    email = html.unescape(email or "").strip().rstrip(".,;:)>]}'\"").lower()
    email = email.replace(" ", "")
    if not email or email.count("@") != 1:
        return None
    if any(email.endswith(suf) for suf in INVALID_SUFFIXES):
        return None
    local, domain = email.split("@", 1)
    if not local or not domain or "." not in domain:
        return None
    if local in INVALID_LOCAL_PARTS:
        return None
    if domain.endswith((".png", ".jpg", ".css", ".js")):
        return None
    # Require a real alphabetic TLD (rejects image names like file@19.59.32)
    if not re.search(r"\.[A-Za-z]{2,}$", domain):
        return None
    if re.fullmatch(r"[\d.]+", domain):
        return None
    tld = domain.rsplit(".", 1)[-1]
    if tld in INVALID_TLDS:
        return None
    # Reject locals that are mostly page-number noise: 235-262@...
    if re.fullmatch(r"[\d\-]+", local):
        return None
    if len(local) > 64 or len(email) > 254:
        return None
    return email


def decode_cfemail(encoded: str) -> str | None:
    try:
        data = bytes.fromhex(encoded)
        if not data:
            return None
        key = data[0]
        decoded = bytes(b ^ key for b in data[1:]).decode("utf-8", errors="ignore")
        return _clean_email(decoded)
    except Exception:
        return None


def extract_emails_from_text(text: str) -> set[str]:
    found: set[str] = set()
    if not text:
        return found

    unescaped = html.unescape(text)

    for match in EMAIL_RE.findall(unescaped):
        email = _clean_email(match)
        if email:
            found.add(email)

    for match in OBFUSCATED_EMAIL_RE.finditer(unescaped):
        local = match.group("local")
        domain_raw = match.group("domain")
        domain = re.sub(
            r"\s*(?:\[dot\]|\(dot\)|\bdot\b)\s*",
            ".",
            domain_raw,
            flags=re.IGNORECASE,
        )
        domain = re.sub(r"\s+", "", domain)
        candidate = f"{local}@{domain}".replace("..", ".")
        email = _clean_email(candidate)
        if email:
            found.add(email)

    for match in CFEMAIL_RE.finditer(unescaped):
        encoded = match.group(1) or match.group(2)
        email = decode_cfemail(encoded)
        if email:
            found.add(email)

    return found


def extract_emails_from_file(path: Path) -> set[str]:
    suffix = path.suffix.lower()
    text = ""
    try:
        if suffix == ".eml":
            from webcrawler.scanner.folder_scanner import _eml_text

            text = _eml_text(path)
        elif suffix in {".html", ".htm", ".txt", ".csv", ".xml"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".pdf":
            text = _pdf_text(path)
        elif suffix == ".docx":
            text = _docx_text(path)
        elif suffix in {".xlsx", ".xlsm"}:
            text = _xlsx_text(path)
    except Exception:
        return set()
    return extract_emails_from_text(text)


def _pdf_text(path: Path) -> str:
    """Extract text from every page of a PDF (author emails can appear late).

    Corrupt or unreadable pages are skipped so one bad page does not abort the file.
    """
    chunks: list[str] = []
    try:
        import fitz

        doc = fitz.open(path)
        try:
            for page in doc:
                try:
                    chunks.append(page.get_text() or "")
                except Exception:
                    continue
        finally:
            doc.close()
        if any(c.strip() for c in chunks):
            return "\n".join(chunks)
    except Exception:
        chunks = []

    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                try:
                    chunks.append(page.extract_text() or "")
                except Exception:
                    continue
    except Exception:
        return ""
    return "\n".join(chunks)


def _docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        try:
            parts.append(paragraph.text or "")
        except Exception:
            continue
    # Tables often hold contact rows; skip broken cells.
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    try:
                        if cell.text:
                            parts.append(cell.text)
                    except Exception:
                        continue
    except Exception:
        pass
    return "\n".join(parts)


def _xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in wb.worksheets:
            try:
                for row in sheet.iter_rows(max_row=500, values_only=True):
                    for cell in row:
                        if cell is not None:
                            parts.append(str(cell))
            except Exception:
                continue
    finally:
        wb.close()
    return "\n".join(parts)
